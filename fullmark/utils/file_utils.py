"""
fullmark/utils/file_utils.py
----------------------------
File type detection, ZIP unpacking, and URL-list extraction utilities.
"""

from __future__ import annotations

import mimetypes
import re
import zipfile
import logging
import tempfile
from pathlib import Path

logger = logging.getLogger(__name__)

# ──────────────────────────────────────────────────────────────────────────────
# Extension routing table
# ──────────────────────────────────────────────────────────────────────────────

DOCUMENT_EXTENSIONS: frozenset[str] = frozenset({
    ".pdf", ".docx", ".doc", ".rtf", ".txt",
    ".epub", ".xlsx", ".xls", ".csv", ".ods",
    ".pptx", ".ppt", ".odp", ".ipynb", ".msg", ".eml",
})

IMAGE_EXTENSIONS: frozenset[str] = frozenset({
    ".jpg", ".jpeg", ".png", ".bmp",
    ".tiff", ".tif", ".webp", ".svg",
})

VIDEO_EXTENSIONS: frozenset[str] = frozenset({
    ".mp4", ".avi", ".mov", ".mkv", ".webm",
    ".mp3", ".wav", ".m4a",
})

WEB_EXTENSIONS: frozenset[str] = frozenset({
    ".html", ".htm",
})

ARCHIVE_EXTENSIONS: frozenset[str] = frozenset({
    ".zip",
})

# Extensions that may contain URL lists — routed to url_list agent
URL_LIST_EXTENSIONS: frozenset[str] = frozenset({
    ".txt", ".docx", ".doc", ".xlsx", ".xls", ".ods", ".csv",
})

# Source-code and config file extensions — routed to document agent (code path)
CODE_EXTENSIONS: frozenset[str] = frozenset({
    # Python
    ".py", ".pyw", ".pyi",
    # JavaScript / TypeScript
    ".js", ".mjs", ".cjs", ".jsx", ".ts", ".tsx",
    # Web front-end
    ".css", ".scss", ".sass", ".less",
    # JVM
    ".java", ".kt", ".kts", ".scala", ".groovy",
    # C family
    ".c", ".h", ".cpp", ".cxx", ".cc", ".hh", ".hpp", ".cs",
    # Systems / low-level
    ".go", ".rs", ".swift", ".zig",
    # Scripting
    ".rb", ".php", ".pl", ".pm", ".lua", ".r", ".rmd",
    # Shell
    ".sh", ".bash", ".zsh", ".fish", ".ps1", ".bat", ".cmd",
    # Functional
    ".hs", ".lhs", ".ml", ".mli", ".ex", ".exs", ".erl", ".clj",
    # Mobile
    ".dart", ".m", ".mm",
    # Data / config
    ".json", ".jsonc", ".json5", ".yaml", ".yml", ".toml",
    ".ini", ".cfg", ".conf",
    # Markup / schema
    ".xml", ".xsd", ".xsl", ".graphql", ".gql", ".proto",
    # Database
    ".sql",
    # Infrastructure
    ".tf", ".tfvars", ".bicep", ".nix",
    # Build systems
    ".gradle", ".cmake", ".makefile",
    # Frameworks
    ".vue", ".svelte",
    # Docs-as-code
    ".rst", ".mdx",
    # Misc text
    ".dockerfile", ".gitignore", ".gitattributes", ".editorconfig",
    ".env.example", ".env.template",
    ".lock",
})

# Maps file extension → fenced-code-block language tag
LANG_MAP: dict[str, str] = {
    ".py": "python", ".pyw": "python", ".pyi": "python",
    ".js": "javascript", ".mjs": "javascript", ".cjs": "javascript", ".jsx": "jsx",
    ".ts": "typescript", ".tsx": "tsx",
    ".css": "css", ".scss": "scss", ".sass": "sass", ".less": "less",
    ".java": "java", ".kt": "kotlin", ".kts": "kotlin",
    ".scala": "scala", ".groovy": "groovy",
    ".c": "c", ".h": "c", ".cpp": "cpp", ".cxx": "cpp",
    ".cc": "cpp", ".hh": "cpp", ".hpp": "cpp",
    ".cs": "csharp", ".go": "go", ".rs": "rust",
    ".swift": "swift", ".zig": "zig", ".dart": "dart",
    ".rb": "ruby", ".php": "php", ".pl": "perl", ".pm": "perl",
    ".lua": "lua", ".r": "r", ".rmd": "r",
    ".sh": "bash", ".bash": "bash", ".zsh": "bash", ".fish": "fish",
    ".ps1": "powershell", ".bat": "batch", ".cmd": "batch",
    ".hs": "haskell", ".lhs": "haskell",
    ".ml": "ocaml", ".mli": "ocaml",
    ".ex": "elixir", ".exs": "elixir",
    ".erl": "erlang", ".hrl": "erlang",
    ".clj": "clojure",
    ".json": "json", ".jsonc": "json", ".json5": "json",
    ".yaml": "yaml", ".yml": "yaml",
    ".toml": "toml", ".ini": "ini", ".cfg": "ini", ".conf": "ini",
    ".xml": "xml", ".xsd": "xml", ".xsl": "xml",
    ".graphql": "graphql", ".gql": "graphql",
    ".proto": "protobuf",
    ".sql": "sql",
    ".tf": "hcl", ".tfvars": "hcl", ".bicep": "bicep",
    ".nix": "nix",
    ".gradle": "groovy", ".cmake": "cmake",
    ".vue": "vue", ".svelte": "svelte",
    ".rst": "rst", ".mdx": "markdown",
    ".dockerfile": "dockerfile",
    ".html": "html", ".htm": "html",
    ".m": "objc", ".mm": "objc",
    ".lock": "yaml",
}

_URL_RE = re.compile(r"^https?://\S+$", re.IGNORECASE)

# GitHub repo URL pattern: https://github.com/{owner}/{repo}[/tree/{branch}[/path]]
_GITHUB_REPO_RE = re.compile(
    r"^https://github\.com/[^/]+/[^/]+"
    r"(?:/tree/[^/]+(?:/.*)?)?$",
    re.IGNORECASE,
)


def detect_agent(source: str | Path) -> str:
    """
    Determine which agent should handle *source*.

    Args:
        source: File path (str or Path) or URL string.

    Returns:
        One of: ``"web"``, ``"document"``, ``"image"``, ``"video"``,
        ``"archive"``, or ``"unknown"``.
    """
    source_str = str(source)

    # GitHub repo URL → repo agent
    if _GITHUB_REPO_RE.match(source_str):
        return "repo"

    # URL → Web Agent
    if source_str.startswith(("http://", "https://")):
        return "web"

    path = Path(source_str)
    ext = path.suffix.lower()

    if ext in ARCHIVE_EXTENSIONS:
        return "archive"
    if ext in DOCUMENT_EXTENSIONS:
        return "document"
    if ext in WEB_EXTENSIONS:
        return "web"
    if ext in IMAGE_EXTENSIONS:
        return "image"
    if ext in VIDEO_EXTENSIONS:
        return "video"
    if ext in CODE_EXTENSIONS:
        return "code"

    # MIME fallback
    mime, _ = mimetypes.guess_type(source_str)
    if mime:
        if mime.startswith("image/"):
            return "image"
        if mime.startswith("video/") or mime.startswith("audio/"):
            return "video"
        if mime in ("text/html", "application/xhtml+xml"):
            return "web"
        if mime in (
            "application/pdf",
            "application/msword",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/vnd.ms-excel",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "text/csv",
            "text/plain",
        ):
            return "document"

    logger.warning("Cannot detect agent for %r — extension %r not in routing table", source_str, ext)
    return "unknown"


def extract_urls_from_file(path: Path) -> tuple[list[str], list[str]]:
    """
    Extract HTTP/HTTPS URLs from a .txt, .docx/.doc, or spreadsheet file.

    Only lines/cells whose stripped content is a valid URL are collected.
    All other lines are recorded as *skipped*.

    Args:
        path: Path to the URL-list file.

    Returns:
        Tuple of (urls, skipped_lines).
        ``urls`` — list of valid URL strings.
        ``skipped_lines`` — list of non-empty lines that were not URLs.
    """
    ext = path.suffix.lower()
    raw_lines: list[str] = []

    if ext == ".txt":
        raw_lines = path.read_text(encoding="utf-8", errors="replace").splitlines()

    elif ext in (".docx", ".doc"):
        try:
            from docx import Document  # type: ignore
            doc = Document(str(path))
            for para in doc.paragraphs:
                raw_lines.append(para.text)
            # Also scan tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        raw_lines.append(cell.text)
        except ImportError:
            logger.warning("python-docx not installed — cannot extract URLs from %s", path.name)
        except Exception as exc:
            logger.warning("Failed to read %s: %s", path.name, exc)

    elif ext in (".xlsx", ".xls", ".ods"):
        try:
            import openpyxl  # type: ignore
            wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    for cell in row:
                        if cell is not None:
                            raw_lines.append(str(cell).strip())
        except ImportError:
            logger.warning("openpyxl not installed — cannot extract URLs from %s", path.name)
        except Exception as exc:
            logger.warning("Failed to read %s: %s", path.name, exc)

    elif ext == ".csv":
        import csv
        try:
            with path.open(newline="", encoding="utf-8", errors="replace") as f:
                for row in csv.reader(f):
                    for cell in row:
                        raw_lines.append(cell.strip())
        except Exception as exc:
            logger.warning("Failed to read %s: %s", path.name, exc)

    urls: list[str] = []
    skipped: list[str] = []
    for line in raw_lines:
        stripped = line.strip()
        if not stripped:
            continue
        if _URL_RE.match(stripped):
            urls.append(stripped)
        else:
            skipped.append(stripped)

    return urls, skipped


def is_url_list_file(path: Path) -> bool:
    """
    Return True if *path* is a supported URL-list file type AND
    contains at least one valid HTTP/HTTPS URL.

    This is a quick heuristic — caller should use ``extract_urls_from_file``
    for the full extraction.
    """
    if path.suffix.lower() not in URL_LIST_EXTENSIONS:
        return False
    try:
        urls, _ = extract_urls_from_file(path)
        return len(urls) > 0
    except Exception:
        return False


def unpack_zip(zip_path: Path) -> tuple[Path, list[Path]]:
    """
    Extract a ZIP archive to a temporary directory.

    Args:
        zip_path: Path to the ZIP file.

    Returns:
        Tuple of (temp_dir_path, list_of_extracted_file_paths).
        Caller is responsible for cleaning up temp_dir_path.
    """
    temp_dir = Path(tempfile.mkdtemp(prefix="fullmark_zip_"))
    extracted: list[Path] = []

    with zipfile.ZipFile(zip_path, "r") as zf:
        for member in zf.infolist():
            if member.is_dir():
                continue
            # Security: strip ALL path components — only use the bare filename.
            # member.filename can contain "../" traversal sequences; using it
            # directly as a dest path would let a crafted ZIP escape temp_dir.
            safe_name = Path(member.filename).name
            if not safe_name:
                continue
            dest = temp_dir / safe_name
            dest.parent.mkdir(parents=True, exist_ok=True)
            dest.write_bytes(zf.read(member))
            extracted.append(dest)
            logger.debug("Extracted %s → %s", member.filename, dest)

    logger.info("Unpacked %s: %d files → %s", zip_path.name, len(extracted), temp_dir)
    return temp_dir, extracted


def url_to_output_name(url: str) -> str:
    """
    Convert a URL to the FullMark naming convention:
    ``<3-letter-domain-prefix>_<path-seg-1>_<path-seg-2>_...``

    Args:
        url: HTTP/HTTPS URL string.

    Returns:
        A safe filename stem (no extension, no path separators).

    Examples::

        https://example.com/docs/api/auth  →  ``exa_docs_api_auth``
        https://docs.python.org/3/library/os.html  →  ``doc_3_library_os``
        https://example.com  →  ``exa``
        https://en.wikipedia.org/wiki/Python  →  ``en__wiki_Python``
    """
    from urllib.parse import urlparse

    parsed = urlparse(url)

    # Domain prefix: first ≤3 alphanumeric chars of netloc (strip www.)
    domain = parsed.netloc.lower()
    if domain.startswith("www."):
        domain = domain[4:]
    domain_clean = re.sub(r"[^a-z0-9]", "", domain)
    domain_prefix = domain_clean[:3] or "url"

    # Path segments: split on /, sanitise each, strip file extension from last
    path = parsed.path.rstrip("/")
    segments: list[str] = []
    if path:
        for seg in path.split("/"):
            if not seg:
                continue
            # Strip file extension from last segment (e.g. .html, .php)
            seg_clean = re.sub(r"\.[a-zA-Z0-9]{1,6}$", "", seg)
            # Keep alphanumerics, hyphens, underscores; truncate to 25 chars
            seg_safe = re.sub(r"[^\w-]", "_", seg_clean)[:25].strip("_")
            if seg_safe:
                segments.append(seg_safe)

    parts = [domain_prefix] + segments
    name = "_".join(parts)
    # Collapse repeated underscores, strip leading/trailing
    name = re.sub(r"_+", "_", name).strip("_")
    return name[:80] or "page"


def safe_output_path(source: str | Path, output_dir: Path, suffix: str = ".md") -> Path:
    """
    Compute a safe output path for *source* inside *output_dir*.

    For URLs applies the FullMark naming convention:
    ``<3-letter-domain>_<path-seg-1>_<path-seg-2>...``

    For file paths uses the stem of the filename.

    If the computed path already exists (collision from a different source with
    the same stem), a short 6-character hash of the full source path is
    appended to disambiguate without overwriting the previous output.

    Args:
        source: Original source path or URL.
        output_dir: Directory to write output into.
        suffix: File suffix for output (default ``.md``).

    Returns:
        Resolved output Path.
    """
    import hashlib as _hashlib

    source_str = str(source)
    if source_str.startswith(("http://", "https://")):
        stem = url_to_output_name(source_str)
    else:
        stem = Path(source_str).stem
        # Sanitise: keep alphanumeric, dash, underscore
        stem = re.sub(r"[^\w-]", "_", stem).strip("_") or "output"

    candidate = output_dir / (stem + suffix)
    if candidate.exists():
        # Append a short hash so a different source with the same stem does not
        # silently overwrite an existing output file.
        short_hash = _hashlib.sha1(source_str.encode()).hexdigest()[:6]
        candidate = output_dir / (stem + "_" + short_hash + suffix)
    return candidate
